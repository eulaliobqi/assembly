"""
Gera artigo.docx a partir de artigo.md (transcriptoma glandula salivar M. spectabilis).
Conversor Markdown -> DOCX (Times New Roman, 12pt, espacamento ~1.5) com insercao de
figuras em linha via marcadores [[FIG:nome_arquivo]] no markdown fonte.

Baseado no padrao de C:\\Users\\eulal\\.claude\\MD-gromacs\\bin\\gerar_artigo_docx.py,
adaptado para inserir figuras em um unico passe (sem pos-processamento por indice de
paragrafo).
"""
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Instalar: pip install python-docx")
    sys.exit(1)

ROOT = Path(__file__).parent.parent
SOURCE = ROOT / "artigo.md"
OUTPUT = ROOT / "artigo.docx"
FIGURES_DIR = ROOT / "figures"

FIGURE_CAPTIONS = {
    "assembly_qc_summary": (
        "Figura 1. Painel didatico de controle de qualidade da montagem e predicao "
        "de proteinas. (A) Numero de sequencias em cada etapa: transcritos "
        "nao-redundantes apos Trinity + CD-HIT-EST a 95% de identidade (103.560, "
        "90.344 genes, GC 34,21%; 12.857 transcritos colapsados por redundancia) e "
        "proteinas preditas por TransDecoder (--single_best_only; 12.445). (B) Curva "
        "Nx dos contigs montados (N10-N50): N50 = 738 pb, ou seja, metade das bases "
        "montadas esta em contigs de pelo menos 738 pb. (C) Completude BUSCO "
        "(insecta_odb10, n=1.367): C:96,2% [S:39,9%, D:56,3%], F:1,9%, M:1,9%. (D) "
        "Distribuicao dos tipos de ORF preditos pelo TransDecoder: completo (7.262; "
        "58,4%), 5'-parcial (2.026; 16,3%), interno (1.986; 16,0%) e 3'-parcial "
        "(1.171; 9,4%)."
    ),
    "annotation_summary": (
        "Figura 2. Cobertura da anotacao funcional e origem taxonomica das 12.445 "
        "proteinas preditas (dados corrigidos pos-fix do bug de merge de anotacao, "
        "Secao 3.2). (A) Proporcao de proteinas anotadas por fonte: dominio Pfam "
        "(62,1%), ortologo KEGG (42,6%), termo GO (44,0%) e hit DIAMOND contra o NCBI "
        "NR (75,1%). (B) Distribuicao taxonomica dos melhores hits DIAMOND: Eukaryota "
        "(8.711; 70,0%), nao classificado (3.164; 25,4%), Bacteria (540; 4,3%), Fungi "
        "(256; 2,1%) e Viruses (30; 0,2%)."
    ),
    "go_distribution": (
        "Figura 3. Distribuicao dos 20 termos GO mais frequentes por categoria entre "
        "as 12.445 proteinas preditas (profundidade GO >= 3). Paineis: Processo "
        "Biologico, Funcao Molecular e Componente Celular, com o numero de proteinas "
        "e a porcentagem associada a cada termo."
    ),
    "kegg_pathways_bar": (
        "Figura 4. Top 20 vias KEGG mais representadas entre as proteinas anotadas, "
        "coloridas por categoria funcional KEGG (Cellular Processes, Environmental/"
        "Genetic Information Processing, Human Diseases, Metabolism, Organismal "
        "Systems, Other). Via mais frequente: Metabolic pathways (845 proteinas, "
        "6,8%), seguida por Biosynthesis of secondary metabolites (388; 3,1%) e "
        "Ribosome (261; 2,1%)."
    ),
    "kegg_pathways_bubble": (
        "Figura 5. Distribuicao das 30 vias KEGG mais representadas, agrupadas por "
        "categoria funcional (eixo x) e numero de proteinas associadas (eixo y; "
        "tamanho do circulo proporcional a contagem). Complementa a Figura 4 ao "
        "evidenciar a dispersao de vias dentro de cada categoria KEGG."
    ),
    "secretome_summary": (
        "Figura 6. Predicao do secretoma classico com TMbed (substitui SignalP/"
        "TMHMM, indisponiveis sem licenca DTU). (A) Funil de predicao: das 12.445 "
        "proteinas, 1.238 (9,9%) tem peptideo sinal, das quais 1.171 (9,4%) "
        "compoem o secretoma classico (<=1 segmento transmembrana) e 67 foram "
        "excluidas por terem mais de 1 segmento TM. (B) Distribuicao do numero de "
        "segmentos transmembrana preditos entre as proteinas com peptideo sinal, "
        "separando secretoma classico e excluidas, com linha de corte em TM=1. "
        "(C) Categorias COG mais frequentes entre as 1.171 proteinas do secretoma "
        "classico: [S] funcao desconhecida (277), [T] transducao de sinal (147), "
        "[O] modificacao pos-traducional/chaperonas (145) e [G] transporte/"
        "metabolismo de carboidratos (89)."
    ),
    "effector_candidates": (
        "Figura 7. Candidatos priorizados a efetor/toxina salivar (n=35), "
        "ranqueados por numero de termos curados correspondentes e por TPM "
        "(cor/tamanho do ponto). Criterio: proteina simultaneamente no secretoma "
        "classico (Figura 6) E com correspondencia a pelo menos um termo/dominio "
        "associado a fitotoxicidade em outros Hemiptera (dominios EF-hand/Ca-"
        "binding: 16; mucina: 9; peptideo/proteina salivar secretado: 4; "
        "fosfolipase A2/B: 2; protease, carboxipeptidase e dipeptidil peptidase 4 "
        "tipo veneno: 1 cada; lacase: 1). Os candidatos mais expressos incluem uma "
        "proteina salivar secretada (melhor hit DIAMOND: Triatoma infestans) e "
        "uma protease tipo veneno (melhor hit: Macrosteles quadrilineatus)."
    ),
}

FIG_MARKER_RE = re.compile(r"^\[\[FIG:([\w-]+)\]\]$")


def set_cell_bg(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def apply_inline(para, text: str):
    """Processa **bold**, *italic* e `code` inline."""
    pattern = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`')
    last = 0
    for m in pattern.finditer(text):
        before = text[last:m.start()]
        if before:
            para.add_run(before)
        if m.group(1) is not None:
            r = para.add_run(m.group(1))
            r.bold = True
        elif m.group(2) is not None:
            r = para.add_run(m.group(2))
            r.italic = True
        elif m.group(3) is not None:
            r = para.add_run(m.group(3))
            r.font.name = 'Courier New'
            r.font.size = Pt(10)
        last = m.end()
    remainder = text[last:]
    if remainder:
        para.add_run(remainder)


def parse_table(lines: list) -> list:
    rows = []
    for line in lines:
        if re.match(r'\|[-: |]+\|', line.strip()):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    return rows


def insert_figure(doc, fig_name: str):
    fig_path = FIGURES_DIR / f"{fig_name}.png"
    caption = FIGURE_CAPTIONS.get(fig_name, f"Figura ({fig_name}).")

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # O estilo 'Normal' usa espacamento de linha EXATO (Pt(18)); sem isto, uma
    # imagem mais alta que 18pt seria cortada pela caixa de linha do paragrafo.
    p_img.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p_img.add_run()
    if fig_path.exists():
        run.add_picture(str(fig_path), width=Inches(6.3))
    else:
        run.text = f"[Figura nao encontrada: {fig_path}]"
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_cap.paragraph_format.space_after = Pt(12)
    r_cap = p_cap.add_run(caption)
    r_cap.italic = True
    r_cap.font.size = Pt(9)


def create_word(md_path: Path, out_path: Path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    para_fmt = style.paragraph_format
    para_fmt.space_after = Pt(6)
    para_fmt.line_spacing = Pt(18)

    for lvl, (name, size, bold) in enumerate([
        ('Heading 1', 16, True),
        ('Heading 2', 14, True),
        ('Heading 3', 13, True),
        ('Heading 4', 12, True),
    ], start=1):
        h = doc.styles[name]
        h.font.name = 'Times New Roman'
        h.font.size = Pt(size)
        h.font.bold = bold
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

    text = md_path.read_text(encoding='utf-8')
    lines = text.splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_heading('', level=1)
            apply_inline(p, line[2:].strip())
            i += 1
            continue

        if line.startswith('## ') and not line.startswith('### '):
            p = doc.add_heading('', level=2)
            apply_inline(p, line[3:].strip())
            i += 1
            continue

        if line.startswith('### ') and not line.startswith('#### '):
            p = doc.add_heading('', level=3)
            apply_inline(p, line[4:].strip())
            i += 1
            continue

        if line.startswith('#### '):
            p = doc.add_heading('', level=4)
            apply_inline(p, line[5:].strip())
            i += 1
            continue

        if line.strip() in ('---', '***', '___'):
            doc.add_paragraph('', style='Normal')
            i += 1
            continue

        if line.strip().startswith('<!--'):
            while i < len(lines) and '-->' not in lines[i]:
                i += 1
            i += 1
            continue

        m_fig = FIG_MARKER_RE.match(line.strip())
        if m_fig:
            insert_figure(doc, m_fig.group(1))
            i += 1
            continue

        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            if not rows:
                continue
            ncols = len(rows[0])
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = 'Table Grid'
            for r_idx, row in enumerate(rows):
                for c_idx, cell_text in enumerate(row):
                    if c_idx >= ncols:
                        continue
                    cell = table.cell(r_idx, c_idx)
                    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                    clean = re.sub(r'\*(.+?)\*', r'\1', clean)
                    clean = re.sub(r'[✅❌⚠️⏳\U0001F4A1]', '', clean).strip()
                    cell.paragraphs[0].clear()
                    run = cell.paragraphs[0].add_run(clean)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    if r_idx == 0:
                        run.bold = True
                        set_cell_bg(cell, 'D9E1F2')
            doc.add_paragraph('', style='Normal')
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip().startswith('- ') or line.strip().startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            content = line.strip()[2:]
            apply_inline(p, content)
            i += 1
            continue

        m_num = re.match(r'^(\d+)\. (.+)', line.strip())
        if m_num:
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            apply_inline(p, m_num.group(2))
            i += 1
            continue

        if line.strip().startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph('\n'.join(code_lines), style='Normal')
                p.runs[0].font.name = 'Courier New'
                p.runs[0].font.size = Pt(9)
                p.paragraph_format.left_indent = Cm(1)
            i += 1
            continue

        if line.strip().startswith('>'):
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.left_indent = Cm(1)
            apply_inline(p, line.strip().lstrip('>').strip())
            for r in p.runs:
                r.italic = True
                r.font.size = Pt(10)
            i += 1
            continue

        p = doc.add_paragraph(style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        apply_inline(p, line.strip())
        i += 1

    doc.save(out_path)
    print(f"[OK] Arquivo salvo: {out_path}")


if __name__ == '__main__':
    create_word(SOURCE, OUTPUT)
